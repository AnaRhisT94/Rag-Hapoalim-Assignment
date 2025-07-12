import os
import json
import pickle
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import asyncio
import hashlib

import pandas as pd
import numpy as np
import torch
import faiss
from docx import Document
from sentence_transformers import SentenceTransformer
from fastapi import FastAPI, HTTPException, BackgroundTasks
from pydantic import BaseModel
import tiktoken
from transformers import AutoTokenizer, AutoModelForCausalLM

from utils import build_chat

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Constants
DATA_DIR = Path("data")
INDEX_DIR = Path("indexes")
CHUNK_SIZE = 2048
TOP_K = 1
EMBEDDING_MODEL_NAME = "facebook/contriever-msmarco"

# Ensure directories exist
DATA_DIR.mkdir(exist_ok=True)
INDEX_DIR.mkdir(exist_ok=True)

# Pydantic models for API
class QueryRequest(BaseModel):
    query: str
    top_k: Optional[int] = TOP_K

class QueryResponse(BaseModel):
    answer: str
    retrieved_chunks: List[str]
    scores: List[float]
    total_chunks_in_index: int

class IndexStatus(BaseModel):
    exists: bool
    total_chunks: int
    last_updated: Optional[str] = None

class VectorStore:
    """Enhanced Vector Store with FAISS backend"""
    
    def __init__(self, dimension: int):
        self.dimension = dimension
        self.index = faiss.IndexFlatIP(dimension)  # Inner product similarity
        self.chunks: List[str] = []
        self.metadata: List[Dict[str, Any]] = []  # Store source info
        
    def add_documents(self, chunks: List[str], embeddings: np.ndarray, metadata: List[Dict[str, Any]]):
        """Add documents with metadata"""
        if len(chunks) != embeddings.shape[0] or len(chunks) != len(metadata):
            raise ValueError("Chunks, embeddings, and metadata must have same length")
            
        # Normalize embeddings for cosine similarity
        faiss.normalize_L2(embeddings.astype(np.float32))
        
        # Add to FAISS index
        self.index.add(embeddings.astype(np.float32))
        
        # Store text and metadata
        self.chunks.extend(chunks)
        self.metadata.extend(metadata)
        
    def search(self, query_embedding: np.ndarray, top_k: int = 3) -> Tuple[List[str], List[float], List[Dict]]:
        """Search with metadata"""
        if len(query_embedding.shape) == 1:
            query_embedding = query_embedding.reshape(1, -1)
            
        # Normalize query embedding
        faiss.normalize_L2(query_embedding.astype(np.float32))
            
        scores, indices = self.index.search(query_embedding.astype(np.float32), top_k)
        
        retrieved_chunks = [self.chunks[idx] for idx in indices[0]]
        retrieved_metadata = [self.metadata[idx] for idx in indices[0]]
        
        return retrieved_chunks, scores[0].tolist(), retrieved_metadata
    
    def save(self, filepath: str):
        """Save vector store to disk"""
        data = {
            'dimension': self.dimension,
            'chunks': self.chunks,
            'metadata': self.metadata
        }
        
        # Save FAISS index
        faiss.write_index(self.index, f"{filepath}.faiss")
        
        # Save other data
        with open(f"{filepath}.pkl", 'wb') as f:
            pickle.dump(data, f)
            
    @classmethod
    def load(cls, filepath: str):
        """Load vector store from disk"""
        # Load FAISS index
        index = faiss.read_index(f"{filepath}.faiss")
        
        # Load other data
        with open(f"{filepath}.pkl", 'rb') as f:
            data = pickle.load(f)
            
        # Reconstruct object
        store = cls(data['dimension'])
        store.index = index
        store.chunks = data['chunks']
        store.metadata = data['metadata']
        
        return store
    
    def get_stats(self) -> Dict[str, Any]:
        """Get statistics about the vector store"""
        return {
            'total_chunks': len(self.chunks),
            'dimension': self.dimension,
            'index_size': self.index.ntotal
        }

class DocumentProcessor:
    """Handle document ingestion from Excel and Word files"""
    
    def __init__(self, tokenizer=None):
        self.tokenizer = tokenizer or tiktoken.get_encoding("cl100k_base")
    
    def load_excel_file(self, filepath: str) -> List[Dict[str, Any]]:
        """Load and process Excel file"""
        try:
            df = pd.read_excel(filepath)
            documents = []
            
            # Find context column (case insensitive)
            context_col = None
            for col in df.columns:
                if 'context' in col.lower():
                    context_col = col
                    break
            
            if context_col is None:
                raise ValueError("No 'context' column found in Excel file")
            
            for idx, row in df.iterrows():
                if pd.notna(row[context_col]):
                    documents.append({
                        'text': str(row[context_col]),
                        'source': 'excel',
                        'source_file': filepath,
                        'row_index': idx,
                        'metadata': row.to_dict()
                    })
            
            logger.info(f"Loaded {len(documents)} documents from Excel file: {filepath}")
            return documents
            
        except Exception as e:
            logger.error(f"Error loading Excel file {filepath}: {e}")
            raise
    
    def load_word_file(self, filepath: str) -> List[Dict[str, Any]]:
        """Load and process Word file"""
        try:
            doc = Document(filepath)
            full_text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            text = '\n'.join(full_text)
            
            documents = [{
                'text': text,
                'source': 'word',
                'source_file': filepath,
                'page_count': len(doc.paragraphs),
                'metadata': {'total_paragraphs': len(doc.paragraphs)}
            }]
            
            logger.info(f"Loaded Word document: {filepath} with {len(doc.paragraphs)} paragraphs")
            return documents
            
        except Exception as e:
            logger.error(f"Error loading Word file {filepath}: {e}")
            raise
    
    def chunk_documents(self, documents: List[Dict[str, Any]], chunk_size: int = CHUNK_SIZE) -> List[Dict[str, Any]]:
        """Split documents into chunks"""
        chunked_docs = []
        
        for doc in documents:
            chunks = split_text(doc['text'], self.tokenizer, chunk_size)
            
            for i, chunk in enumerate(chunks):
                chunk_metadata = doc['metadata'].copy() if 'metadata' in doc else {}
                chunk_metadata.update({
                    'chunk_index': i,
                    'total_chunks': len(chunks),
                    'source': doc['source'],
                    'source_file': doc['source_file']
                })
                
                chunked_docs.append({
                    'text': chunk,
                    'metadata': chunk_metadata
                })
        
        logger.info(f"Created {len(chunked_docs)} chunks from {len(documents)} documents")
        return chunked_docs

class RAGService:
    """Main RAG service combining all components"""
    
    def __init__(self, model_path: str = 'togethercomputer/Llama-2-7B-32K-Instruct', model_name: str = 'llama2'):
        self.embedding_model = SentenceTransformer(EMBEDDING_MODEL_NAME)
        self.embedding_dim = self.embedding_model.get_sentence_embedding_dimension()
        
        self.vector_store: Optional[VectorStore] = None
        self.processor = DocumentProcessor()
        
        # LLM components
        self.model_path = model_path
        self.model_name = model_name
        self.qa_model = None
        self.qa_tokenizer = None
        
        if model_path:
            self._load_qa_model()
    
    def _load_qa_model(self):
        """Load the QA model and tokenizer"""
        try:
            if self.qa_tokenizer is None:
                self.qa_tokenizer = AutoTokenizer.from_pretrained(self.model_path, trust_remote_code=True)
            if self.qa_model is None:
                self.qa_model = AutoModelForCausalLM.from_pretrained(
                    self.model_path,
                    torch_dtype=torch.float16,
                    device_map="cuda:0"
                )
            logger.info(f"Loaded QA model: {self.model_path}")
        except Exception as e:
            logger.error(f"Error loading QA model: {e}")
            raise
    
    def _get_index_hash(self) -> str:
        """Generate hash of data files for cache invalidation"""
        hasher = hashlib.md5()
        
        for file_path in DATA_DIR.glob("*"):
            if file_path.is_file() and file_path.suffix in ['.xlsx', '.xls', '.docx']:
                hasher.update(str(file_path.stat().st_mtime).encode())
                hasher.update(file_path.name.encode())
        
        return hasher.hexdigest()
    
    def ingestion(self) -> List[Dict[str, Any]]:
        """Load documents from data directory"""
        documents = []
        
        # Process Excel files
        for excel_file in DATA_DIR.glob("*.xlsx"):
            documents.extend(self.processor.load_excel_file(str(excel_file)))
        
        for excel_file in DATA_DIR.glob("*.xls"):
            documents.extend(self.processor.load_excel_file(str(excel_file)))
        
        # Process Word files
        for word_file in DATA_DIR.glob("*.docx"):
            documents.extend(self.processor.load_word_file(str(word_file)))
        
        if not documents:
            raise ValueError("No documents found in data directory")
        
        return documents
    
    def indexing(self, force_rebuild: bool = False) -> bool:
        """Create or load FAISS index"""
        index_path = INDEX_DIR / "vector_store"
        hash_path = INDEX_DIR / "data_hash.txt"
        
        current_hash = self._get_index_hash()
        
        # Check if we can use existing index
        if not force_rebuild and index_path.with_suffix('.faiss').exists():
            try:
                # Check if hash matches
                if hash_path.exists():
                    with open(hash_path, 'r') as f:
                        stored_hash = f.read().strip()
                    
                    if stored_hash == current_hash:
                        logger.info("Loading existing index...")
                        self.vector_store = VectorStore.load(str(index_path))
                        return False  # Didn't rebuild
                
                logger.info("Data files changed, rebuilding index...")
            except Exception as e:
                logger.warning(f"Error loading existing index: {e}. Rebuilding...")
        
        # Build new index
        logger.info("Building new index...")
        
        # Ingestion
        documents = self.ingestion()
        
        # Chunking
        chunked_docs = self.processor.chunk_documents(documents, CHUNK_SIZE)
        
        # Extract texts and metadata
        texts = [doc['text'] for doc in chunked_docs]
        metadata = [doc['metadata'] for doc in chunked_docs]
        
        # Generate embeddings
        logger.info("Generating embeddings...")
        embeddings = self.embedding_model.encode(texts, show_progress_bar=True)
        
        # Create vector store
        self.vector_store = VectorStore(self.embedding_dim)
        self.vector_store.add_documents(texts, embeddings, metadata)
        
        # Save index and hash
        self.vector_store.save(str(index_path))
        with open(hash_path, 'w') as f:
            f.write(current_hash)
        
        logger.info(f"Index built successfully with {len(texts)} chunks")
        return True  # Rebuilt
    
    def retrieval(self, query: str, top_k: int = TOP_K) -> Tuple[List[str], List[float], List[Dict]]:
        """Retrieve relevant chunks"""
        if self.vector_store is None:
            raise ValueError("Vector store not initialized. Run indexing first.")
        
        # Generate query embedding
        query_embedding = self.embedding_model.encode([query])
        
        # Search
        chunks, scores, metadata = self.vector_store.search(query_embedding, top_k)
        
        return chunks, scores, metadata
    
    def generation(self, query: str, context_chunks: List[str]) -> str:
        """Generate answer using retrieved context"""
        if self.qa_model is None:
            raise ValueError("QA model not loaded")
        
        # Build context
        context = "\n\n".join([f"{i}: {chunk}" for i, chunk in enumerate(context_chunks)])
        
        prompt = f"Context:\n{context}\n\nQuestion: {query}\n\nAnswer:"
        
        # Format for model
        formatted_prompt = build_chat(self.qa_tokenizer, prompt, self.model_name)
        
        # Generate
        inputs = self.qa_tokenizer(formatted_prompt, return_tensors="pt", truncation=True, max_length=8192)
        inputs = {k: v.to(self.qa_model.device) for k, v in inputs.items()}
        
        with torch.inference_mode():
            outputs = self.qa_model.generate(
                **inputs,
                max_new_tokens=16,
                do_sample=False,
                num_beams=1,
                temperature=1.0,
                pad_token_id=self.qa_tokenizer.eos_token_id
            )
        
        # Decode response
        response = self.qa_tokenizer.decode(
            outputs[0][inputs['input_ids'].shape[1]:], 
            skip_special_tokens=True
        )
        
        del outputs
        torch.cuda.empty_cache()
        
        return response.strip()
    
    def get_index_status(self) -> IndexStatus:
        """Get current index status"""
        index_path = INDEX_DIR / "vector_store.faiss"
        
        if index_path.exists() and self.vector_store:
            stats = self.vector_store.get_stats()
            return IndexStatus(
                exists=True,
                total_chunks=stats['total_chunks'],
                last_updated=str(index_path.stat().st_mtime)
            )
        else:
            return IndexStatus(exists=False, total_chunks=0)

# Initialize service
# rag_service = RAGService(model_path='togethercomputer/Llama-2-7B-32K-Instruct', model_name='llama2') # Removed global initialization

# Create a singleton pattern for the RAG service
rag_service = None

def get_rag_service():
    global rag_service
    if rag_service is None:
        rag_service = RAGService(
            model_path='togethercomputer/Llama-2-7B-32K-Instruct',
            model_name='llama2'
        )
    return rag_service

# FastAPI app
app = FastAPI(title="RAG Service API", version="1.0.0")

@app.on_event("startup")
async def startup_event():
    """Initialize the service on startup"""
    try:
        logger.info("Starting RAG service...")
        service = get_rag_service()
        service.indexing()  # This will load existing or create new index
        logger.info("RAG service started successfully")
    except Exception as e:
        logger.error(f"Failed to start RAG service: {e}")

@app.post("/query", response_model=QueryResponse)
async def query_endpoint(request: QueryRequest):
    """Main query endpoint"""
    try:
        service = get_rag_service()
        # Retrieval
        chunks, scores, metadata = service.retrieval(request.query, request.top_k)
        
        # Generation (if QA model is available)
        if service.qa_model:
            answer = service.generation(request.query, chunks)
        else:
            answer = "QA model not available. Only retrieval performed."
        
        return QueryResponse(
            answer=answer,
            retrieved_chunks=chunks,
            scores=scores,
            total_chunks_in_index=len(service.vector_store.chunks)
        )
        
    except Exception as e:
        logger.error(f"Error processing query: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/rebuild-index")
async def rebuild_index(background_tasks: BackgroundTasks):
    """Force rebuild the index"""
    try:
        service = get_rag_service()
        background_tasks.add_task(service.indexing, force_rebuild=True)
        return {"message": "Index rebuild started in background"}
    except Exception as e:
        logger.error(f"Error rebuilding index: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/index-status", response_model=IndexStatus)
async def get_index_status():
    """Get current index status"""
    service = get_rag_service()
    return service.get_index_status()

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    service = get_rag_service()
    return {
        "status": "healthy",
        "index_exists": service.vector_store is not None,
        "qa_model_loaded": service.qa_model is not None
    }

if __name__ == "__main__":
    import uvicorn
    
    uvicorn.run(
        "rag_service:app",
        host="0.0.0.0",
        port=8000,
        reload=False,  
        workers=1,     
        log_level="debug"
    ) 